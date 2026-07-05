#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════╗
║         ALT KILL ZONE  v3.1             ║
║         AUTHOR: PARADOX MX              ║
║  OSINT | DNS | ARCHIVOS | APK | WiFi    ║
║  USERNAME SEARCH | REPORTE WORD         ║
╚══════════════════════════════════════════╝

"""

import sys
import os
import re
import time
import json
import socket
import platform
import hashlib
import threading
import subprocess
import argparse
import webbrowser
import random
from datetime import datetime
from urllib.parse import quote, urlencode
from concurrent.futures import ThreadPoolExecutor, as_completed

# ── Imports opcionales ────────────────────────────────────────
try:
    import requests
    REQUESTS_OK = True
except ImportError:
    REQUESTS_OK = False

try:
    import phonenumbers
    from phonenumbers import carrier, geocoder, timezone as ph_timezone
    PHONE_OK = True
except ImportError:
    PHONE_OK = False

try:
    import whois as whois_lib
    WHOIS_LIB_OK = True
except ImportError:
    WHOIS_LIB_OK = False

try:
    from bs4 import BeautifulSoup
    BS4_OK = True
except ImportError:
    BS4_OK = False

try:
    from docx import Document
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ══════════════════════════════════════════
#   DETECCIÓN DE PLATAFORMA
# ══════════════════════════════════════════
SISTEMA = platform.system()  # 'Linux', 'Windows', 'Darwin'
ES_TERMUX = "com.termux" in os.environ.get("PREFIX", "") or os.path.isdir("/data/data/com.termux")
ES_WINDOWS = SISTEMA == "Windows"
ES_LINUX = SISTEMA == "Linux" and not ES_TERMUX
ES_MAC = SISTEMA == "Darwin"

if ES_WINDOWS:
    try:
        os.system("")
    except Exception:
        pass

_FORZAR_SIN_COLOR = os.environ.get("NO_COLOR") is not None or not sys.stdout.isatty()

# ══════════════════════════════════════════
#   COLORES
# ══════════════════════════════════════════
if _FORZAR_SIN_COLOR:
    R = G = Y = C = W = B = RE = ""
else:
    R  = "\033[1;31m"
    G  = "\033[1;32m"
    Y  = "\033[1;33m"
    C  = "\033[1;36m"
    W  = "\033[1;37m"
    B  = "\033[0;90m"
    RE = "\033[0m"

# ══════════════════════════════════════════
#   SESIÓN (para el reporte Word consolidado)
# ══════════════════════════════════════════
SESION_RESULTADOS = []  # [{"tipo","objetivo","fecha","datos"}, ...]

def _agregar_sesion(tipo, objetivo, datos):
    if datos is None:
        return
    SESION_RESULTADOS.append({
        "tipo": tipo,
        "objetivo": objetivo,
        "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "datos": datos,
    })

# ══════════════════════════════════════════
#   BANNER
# ══════════════════════════════════════════
def limpiar_pantalla():
    try:
        os.system("cls" if ES_WINDOWS else "clear")
    except Exception:
        pass

_SIMBOLOS_OCULTOS = ["⛧", "☠", "▲", "☾", "⚚", "☥", "✦"]

def show_banner():
    limpiar_pantalla()
    simbolo = random.choice(_SIMBOLOS_OCULTOS)
    print(f"""{R}
   {simbolo}═══════════════════════════════════════════════════{simbolo}
   ░█████╗░██╗░░░░░████████╗  ██╗░░██╗██╗██╗░░░░░██╗░░░░░
   ██╔══██╗██║░░░░░╚══██╔══╝  ██║░██╔╝██║██║░░░░░██║░░░░░
   ███████║██║░░░░░░░░██║░░░  █████╔╝░██║██║░░░░░██║░░░░░
   ██╔══██║██║░░░░░░░░██║░░░  ██╔═██╗░██║██║░░░░░██║░░░░░
   ██║░░██║███████╗░░░██║░░░  ██║░╚██╗██║███████╗███████╗
   ╚═╝░░╚═╝╚══════╝░░░╚═╝░░░  ╚═╝░░╚═╝╚═╝╚══════╝╚══════╝
{Y}   ░░░░░░░░░█████╗░░█████╗░███╗░░██╗███████╗░░░░░░░░░
   ░░░░░░░░░╚════██╗██╔══██╗████╗░██║██╔════╝░░░░░░░░░
   ░░░░░░░░░░░███╔═╝██║░░██║██╔██╗██║█████╗░░░░░░░░░░░
   ░░░░░░░░░██╔══╝░░██║░░██║██║╚████║██╔══╝░░░░░░░░░░░
   ░░░░░░░░░███████╗╚█████╔╝██║░╚███║███████╗░░░░░░░░░
   ░░░░░░░░░╚══════╝░╚════╝░╚═╝░░╚══╝╚══════╝░░░░░░░░░
{R}                    {simbolo}      ▲      {simbolo}
                       ╱   ╲
                      ╱  ◉  ╲
                     ╱_______╲
{W}         MULTI-TOOL v3.1  |  AUTHOR: PARADOX MX
{B}   [OSINT][DNS][ARCHIVOS][APK][WiFi][USERNAME][WORD]
   Plataforma detectada: {_nombre_plataforma()}
{R}   {simbolo}═══════════════════════════════════════════════════{simbolo}{RE}
""")

def _nombre_distro_linux():
    """Lee /etc/os-release para obtener el nombre de la distribución Linux."""
    try:
        info = {}
        with open("/etc/os-release", "r", encoding="utf-8") as f:
            for linea in f:
                if "=" in linea:
                    clave, _, valor = linea.strip().partition("=")
                    info[clave] = valor.strip('"')
        nombre = info.get("PRETTY_NAME") or info.get("NAME")
        if nombre:
            return nombre
    except Exception:
        pass

    # Fallback: intentar con lsb_release si /etc/os-release no existe
    ok, out = _correr_comando(["lsb_release", "-ds"])
    if ok and out.strip():
        return out.strip().strip('"')

    return None

def _nombre_plataforma():
    if ES_TERMUX:
        return "Termux (Android)"
    if ES_WINDOWS:
        return "Windows"
    if ES_MAC:
        return "macOS"
    if ES_LINUX:
        distro = _nombre_distro_linux()
        if distro:
            return f"Linux — {distro}"
        return "Linux (distribución no identificada)"
    return SISTEMA or "Desconocida"

# ══════════════════════════════════════════
#   UTILIDADES
# ══════════════════════════════════════════
def log(tipo, msg):
    iconos = {
        "info": f"{C}[*]{RE}",
        "ok"  : f"{G}[+]{RE}",
        "err" : f"{R}[-]{RE}",
        "warn": f"{Y}[!]{RE}"
    }
    print(f"{iconos.get(tipo, '[?]')} {msg}")

def separador(titulo=""):
    linea = f"{B}{'═'*50}{RE}"
    if titulo:
        print(f"\n{linea}")
        print(f"{Y}  {titulo}{RE}")
        print(f"{linea}")
    else:
        print(linea)

def _directorio_salida():
    home = os.path.expanduser("~")
    candidatos = [
        os.path.join(home, "storage", "downloads"),
        os.path.join(home, "Downloads"),
        home,
    ]
    for d in candidatos:
        if os.path.isdir(d):
            return d
    return os.getcwd()

def guardar(nombre, datos: dict):
    ts  = datetime.now().strftime("%Y%m%d_%H%M%S")
    d   = _directorio_salida()
    ruta = os.path.join(d, f"akz_{nombre}_{ts}.json")
    try:
        with open(ruta, "w", encoding="utf-8") as f:
            json.dump(datos, f, indent=2, ensure_ascii=False)
        log("ok", f"Guardado → {ruta}")
        return ruta
    except Exception as e:
        ruta = f"akz_{nombre}_{ts}.json"
        try:
            with open(ruta, "w", encoding="utf-8") as f:
                json.dump(datos, f, indent=2, ensure_ascii=False)
            log("ok", f"Guardado → {ruta}")
            return ruta
        except Exception as e2:
            log("err", f"No se pudo guardar el reporte: {e2}")
            return None

def _correr_comando(cmd, timeout=10):
    try:
        out = subprocess.check_output(
            cmd, stderr=subprocess.DEVNULL, text=True, timeout=timeout
        )
        return True, out
    except FileNotFoundError:
        return False, ""
    except Exception:
        return False, ""

# ══════════════════════════════════════════
#   MÓDULO 1 — OSINT IP
# ══════════════════════════════════════════
def decimal_a_dms(lat, lon):
    def convertir(grados):
        g = int(abs(grados))
        m = int((abs(grados) - g) * 60)
        s = round(((abs(grados) - g) * 60 - m) * 60, 2)
        return g, m, s

    lat_g, lat_m, lat_s = convertir(lat)
    lon_g, lon_m, lon_s = convertir(lon)
    dir_lat = "N" if lat >= 0 else "S"
    dir_lon = "E" if lon >= 0 else "W"
    lat_dms = f"{lat_g}°{lat_m}'{lat_s}\"{dir_lat}"
    lon_dms = f"{lon_g}°{lon_m}'{lon_s}\"{dir_lon}"
    return lat_dms, lon_dms

def osint_ip(ip):
    separador("OSINT · IP")
    if not REQUESTS_OK:
        log("err", "Falta la librería requests → pip install requests"); return None
    resultado = {"ip": ip}
    try:
        res = requests.get(f"https://ipwho.is/{ip}", timeout=8).json()

        if not res.get("success", False):
            log("err", f"IP inválida o privada: {res.get('message', '')}"); return None

        lat = res.get("latitude")
        lon = res.get("longitude")

        campos = {
            "IP"           : res.get("ip"),
            "Tipo"         : res.get("type"),
            "País"         : res.get("country"),
            "Código"       : res.get("country_code"),
            "Región"       : res.get("region"),
            "Ciudad"       : res.get("city"),
            "Código postal": res.get("postal"),
            "Latitud"      : lat,
            "Longitud"     : lon,
            "Zona horaria" : res.get("timezone", {}).get("id", "N/A"),
            "UTC offset"   : res.get("timezone", {}).get("utc", "N/A"),
            "Proveedor"    : res.get("connection", {}).get("isp", "N/A"),
            "Organización" : res.get("connection", {}).get("org", "N/A"),
            "ASN"          : res.get("connection", {}).get("asn", "N/A"),
        }
        for k, v in campos.items():
            log("ok", f"{k:14}: {W}{v}{RE}")
            resultado[k] = v

        if lat is not None and lon is not None:
            lat_dms, lon_dms = decimal_a_dms(lat, lon)
            lat_enc = quote(lat_dms)
            lon_enc = quote(lon_dms)
            link_decimal = f"https://www.google.com/maps/@{lat},{lon},8z"
            link_dms     = f"https://www.google.com/maps/place/{lat_enc}+{lon_enc}/"
            print()
            log("ok", f"{'Coords DMS':14}: {W}{lat_dms}, {lon_dms}{RE}")
            log("ok", f"{'Mapa decimal':14}: {C}{link_decimal}{RE}")
            log("ok", f"{'Mapa DMS':14}: {C}{link_dms}{RE}")
            resultado["coords_dms"]   = f"{lat_dms}, {lon_dms}"
            resultado["mapa_decimal"] = link_decimal
            resultado["mapa_dms"]     = link_dms
        else:
            log("warn", "Coordenadas no disponibles para esta IP")

    except requests.exceptions.RequestException as e:
        log("err", f"Error de red: {e}")
    except Exception as e:
        log("err", str(e))
    guardar("osint_ip", resultado)
    return resultado

# ══════════════════════════════════════════
#   MÓDULO 2 — OSINT EMAIL
# ══════════════════════════════════════════
class OsintEmail:


    def __init__(self, email):
        self.email = email
        self.investigacion = {
            "email"          : email,
            "fecha_analisis" : datetime.now().isoformat(),
            "redes_sociales" : {},
            "info_dominio"   : {},
            "reputacion"     : {},
            "dorks"          : [],
            "enlaces_breach" : [],
        }

    def _validar(self):
        return re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', self.email)

    def _partes(self):
        if "@" in self.email:
            u, d = self.email.split("@", 1)
            return u, d
        return "", ""

    def _patron_usuario(self, u):
        if not u:
            return "No identificable"
        if "." in u:
            return "nombre.apellido (estilo profesional)"
        if "_" in u:
            return "nombre_apellido (estilo técnico)"
        if any(c.isdigit() for c in u):
            if u[-2:].isdigit():
                return "nombre+año (posible año de nacimiento)"
            return "nombre+números"
        if len(u) <= 3:
            return "iniciales / abreviación"
        if len(u) >= 12:
            return "usuario largo / posiblemente generado"
        return "usuario simple"

    def verificar_filtraciones(self):
        separador("EMAIL · VERIFICACIÓN DE FILTRACIONES")
        log("warn", "Este programa NO consulta bases de datos de filtraciones")
        log("warn", "directamente (la API de HIBP requiere suscripción paga).")
        log("info", "Usa estos verificadores oficiales con tu propio email:\n")
        enlaces = [
            ("Have I Been Pwned", f"https://haveibeenpwned.com/account/{quote(self.email)}"),
            ("Firefox Monitor",   "https://monitor.firefox.com/"),
            ("DeHashed (manual)", "https://www.dehashed.com/"),
        ]
        for nombre, url in enlaces:
            print(f"  {C}→ {nombre:20}{RE} {url}")
            self.investigacion["enlaces_breach"].append({"servicio": nombre, "url": url})

    def buscar_redes_sociales(self):
        separador("EMAIL · POSIBLES PERFILES (a verificar manualmente)")
        usuario, _ = self._partes()
        plataformas = [
            ("GitHub",    f"https://github.com/{usuario}"),
            ("Twitter/X", f"https://twitter.com/{usuario}"),
            ("Instagram", f"https://instagram.com/{usuario}"),
            ("LinkedIn",  f"https://linkedin.com/in/{usuario}"),
            ("Reddit",    f"https://reddit.com/user/{usuario}"),
            ("TikTok",    f"https://tiktok.com/@{usuario}"),
            ("YouTube",   f"https://youtube.com/@{usuario}"),
            ("Twitch",    f"https://twitch.tv/{usuario}"),
        ]
        log("info", "Estas son URLs candidatas basadas en el usuario del email.")
        log("info", "Usa el módulo de búsqueda de username para verificar cuáles existen realmente.\n")
        for nombre, url in plataformas:
            print(f"  {B}{nombre:<12}{RE} {C}{url}{RE}")
            self.investigacion["redes_sociales"][nombre] = url

    def analizar_reputacion(self):
        separador("EMAIL · REPUTACIÓN DEL DOMINIO")
        usuario, dominio = self._partes()
        rep_db = {
            "gmail.com"      : ("Alta",     "Personal"),
            "outlook.com"    : ("Alta",     "Personal"),
            "yahoo.com"      : ("Media",    "Personal"),
            "hotmail.com"    : ("Media",    "Personal"),
            "protonmail.com" : ("Muy alta", "Privado / cifrado"),
            "icloud.com"     : ("Alta",     "Personal"),
            "aol.com"        : ("Baja",     "Legacy"),
        }
        rep, tipo = rep_db.get(dominio, ("Desconocida", "No identificado / posible corporativo"))
        patron    = self._patron_usuario(usuario)
        log("ok", f"{'Dominio':16}: {W}{dominio}{RE}")
        log("ok", f"{'Reputación':16}: {W}{rep}{RE}")
        log("ok", f"{'Tipo cuenta':16}: {W}{tipo}{RE}")
        log("ok", f"{'Patrón usuario':16}: {W}{patron}{RE}")
        self.investigacion["reputacion"] = {
            "dominio": dominio, "nivel_reputacion": rep,
            "tipo_cuenta": tipo, "patron_usuario": patron
        }

    def analisis_dominio(self):
        separador("EMAIL · INFO TÉCNICA DEL DOMINIO")
        usuario, dominio = self._partes()
        log("info", f"Dominio objetivo: {dominio}")
        try:
            ip = socket.gethostbyname(dominio)
            log("ok", f"IP dominio     : {W}{ip}{RE}")
            self.investigacion["info_dominio"]["ip"] = ip
        except Exception:
            log("warn", "No se pudo resolver el dominio")

        ok, out = _correr_comando(["nslookup", "-type=MX", dominio])
        if ok and "mail exchanger" in out.lower():
            log("ok", f"Registros MX   : {G}encontrados{RE}")
            self.investigacion["info_dominio"]["tiene_mx"] = True
        elif ok:
            log("warn", "Sin registros MX detectados")
            self.investigacion["info_dominio"]["tiene_mx"] = False
        else:
            log("warn", "nslookup no disponible en este sistema")

        for srv in [f"smtp.{dominio}", f"imap.{dominio}", f"pop3.{dominio}"]:
            try:
                ip2 = socket.gethostbyname(srv)
                log("ok", f"{srv:30} → {W}{ip2}{RE}")
            except Exception:
                pass

        self.investigacion["info_dominio"]["dominio"] = dominio
        self.investigacion["info_dominio"]["tipo"] = (
            "Email gratuito" if dominio in ["gmail.com", "yahoo.com", "hotmail.com", "outlook.com"]
            else "Posible email corporativo / propio"
        )

    def busquedas_avanzadas(self):
        separador("EMAIL · GOOGLE DORKS")
        usuario, dominio = self._partes()
        dorks = [
            f'"{self.email}"',
            f'intext:"{self.email}" filetype:pdf',
            f'site:linkedin.com "{self.email}"',
            f'site:github.com "{usuario}"',
            f'site:pastebin.com "{self.email}"',
            f'filetype:xls "{self.email}"',
            f'intitle:"{usuario}"',
            f'inurl:"{usuario}"',
        ]
        for i, dork in enumerate(dorks, 1):
            url = f"https://www.google.com/search?{urlencode({'q': dork})}"
            print(f"  {Y}[{i}]{RE} {dork}")
            print(f"      {C}→ {url}{RE}")
            self.investigacion["dorks"].append({"dork": dork, "url": url})

    def analisis_completo(self):
        separador("EMAIL · ANÁLISIS COMPLETO")
        log("info", f"Objetivo: {W}{self.email}{RE}")
        log("info", f"Inicio  : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print()
        self.verificar_filtraciones()
        self.analizar_reputacion()
        self.buscar_redes_sociales()
        self.analisis_dominio()
        self.busquedas_avanzadas()
        print()
        log("ok", f"{G}ANÁLISIS COMPLETO FINALIZADO{RE}")

    def generar_reporte(self):
        separador("EMAIL · GENERAR REPORTE")
        usuario, dominio = self._partes()
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        ruta = guardar(f"email_{usuario}_{ts}", self.investigacion)
        if ruta:
            log("ok", f"Reporte JSON disponible en: {ruta}")

    def menu(self):
        if not self._validar():
            log("err", f"Email con formato inválido: {self.email}"); return
        usuario, dominio = self._partes()
        separador(f"EMAIL OSINT → {self.email}")
        log("ok", f"Usuario : {W}{usuario}{RE}")
        log("ok", f"Dominio : {W}{dominio}{RE}")
        while True:
            print(f"""
{R}┌────────────────────────────────────────────────┐
│          ⚙  MENÚ EMAIL OSINT                   │
├────────────────────────────────────────────────┤
│  {G}[1]{W}  🔓 Verificadores de filtraciones (enlaces) {R}│
│  {G}[2]{W}  🌐 Posibles perfiles en redes sociales    {R}│
│  {G}[3]{W}  📊 Reputación del dominio                 {R}│
│  {G}[4]{W}  🔧 Análisis técnico del dominio            {R}│
│  {G}[5]{W}  🔍 Google Dorks                            {R}│
│  {G}[6]{W}  🚀 Análisis completo                       {R}│
│  {G}[7]{W}  📄 Generar reporte                         {R}│
│  {G}[0]{W}  ❌ Volver                                  {R}│
└────────────────────────────────────────────────┘{RE}""")
            op = input(f"  {Y}[?]{RE} Opción (0-7): ").strip()
            if   op == "1": self.verificar_filtraciones()
            elif op == "2": self.buscar_redes_sociales()
            elif op == "3": self.analizar_reputacion()
            elif op == "4": self.analisis_dominio()
            elif op == "5": self.busquedas_avanzadas()
            elif op == "6": self.analisis_completo()
            elif op == "7": self.generar_reporte()
            elif op == "0": break
            else: log("err", "Opción inválida")
            input(f"\n  {C}[↵] Presiona Enter para continuar...{RE}")


def osint_email(email):
    obj = OsintEmail(email)
    obj.menu()
    return obj.investigacion

# ══════════════════════════════════════════
#   MÓDULO 3 — OSINT TELÉFONO
# ══════════════════════════════════════════
_PREFIJOS = {
    "1":"Estados Unidos / Canadá","52":"México","57":"Colombia",
    "54":"Argentina","55":"Brasil","56":"Chile","51":"Perú",
    "58":"Venezuela","593":"Ecuador","595":"Paraguay","598":"Uruguay",
    "591":"Bolivia","502":"Guatemala","503":"El Salvador","504":"Honduras",
    "505":"Nicaragua","506":"Costa Rica","507":"Panamá",
    "1809":"República Dominicana","1787":"Puerto Rico",
    "34":"España","33":"Francia","44":"Reino Unido","49":"Alemania",
    "39":"Italia","7":"Rusia","86":"China","81":"Japón","82":"Corea del Sur",
    "91":"India","61":"Australia","64":"Nueva Zelanda",
}

_TIPOS = {
    0:"Fijo",1:"Móvil",2:"Sin cargo (800)",3:"Premium",
    4:"Compartido",5:"VoIP / Internet",6:"Paginador",
    7:"UAN",8:"Desconocido",27:"Móvil o fijo",
}

def _tipo_numero(parsed):
    t = phonenumbers.number_type(parsed)
    return _TIPOS.get(int(str(t).split(".")[-1]), str(t))

def osint_phone(numero):
    separador("OSINT · TELÉFONO")
    if not PHONE_OK:
        log("err", "Falta la librería phonenumbers → pip install phonenumbers"); return None

    if not numero.startswith("+"):
        numero = "+" + numero

    try:
        parsed = phonenumbers.parse(numero, None)
    except Exception as e:
        log("err", f"No se pudo parsear el número: {e}"); return None

    es_valido  = phonenumbers.is_valid_number(parsed)
    es_posible = phonenumbers.is_possible_number(parsed)

    if not es_posible:
        log("err", "Número imposible (longitud incorrecta)"); return None
    if not es_valido:
        log("warn", "Número posible pero no confirmado como válido")

    operador  = carrier.name_for_number(parsed, "es") or carrier.name_for_number(parsed, "en") or "N/A"
    region    = geocoder.description_for_number(parsed, "es") or "N/A"
    zonas     = list(ph_timezone.time_zones_for_number(parsed))
    pais_code = phonenumbers.region_code_for_number(parsed)
    tipo      = _tipo_numero(parsed)
    fmt_intl  = phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.INTERNATIONAL)
    fmt_e164  = phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164)
    fmt_nac   = phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.NATIONAL)
    cod_pais  = parsed.country_code

    pais_nombre = "N/A"
    for prefijo in sorted(_PREFIJOS.keys(), key=len, reverse=True):
        if fmt_e164.lstrip("+").startswith(prefijo):
            pais_nombre = _PREFIJOS[prefijo]; break

    print(f"""
{R}┌──────────────────────────────────────────────────┐
│         📱 RESULTADO OSINT TELÉFONO              │
├──────────────────────────────────────────────────┤{RE}
  {Y}Número ingresado{RE}: {W}{numero}{RE}
  {Y}Formato E.164   {RE}: {W}{fmt_e164}{RE}
  {Y}Formato intl.   {RE}: {W}{fmt_intl}{RE}
  {Y}Formato nac.    {RE}: {W}{fmt_nac}{RE}
{R}├──────────────────────────────────────────────────┤{RE}
  {Y}Código de país  {RE}: {W}+{cod_pais}{RE}
  {Y}País (código)   {RE}: {W}{pais_code}{RE}
  {Y}País (nombre)   {RE}: {W}{pais_nombre}{RE}
  {Y}Región          {RE}: {W}{region}{RE}
  {Y}Operador        {RE}: {W}{operador}{RE}
  {Y}Tipo de línea   {RE}: {W}{tipo}{RE}
  {Y}Zona(s) horaria {RE}: {W}{", ".join(zonas) if zonas else "N/A"}{RE}
  {Y}Número válido   {RE}: {G if es_valido else R}{es_valido}{RE}
{R}├──────────────────────────────────────────────────┤{RE}""")

    dorks = [
        f'"{fmt_intl}"',
        f'"{fmt_nac}"',
        f'"{numero}"',
        f'site:linkedin.com "{fmt_intl}"',
    ]
    print(f"  {C}🔍 Google Dorks:{RE}")
    for d in dorks:
        url = f"https://www.google.com/search?{urlencode({'q': d})}"
        print(f"    {B}→{RE} {d}")
        print(f"       {C}{url}{RE}")
    print(f"{R}└──────────────────────────────────────────────────┘{RE}")

    resultado = {
        "numero_ingresado": numero, "e164": fmt_e164, "formato_intl": fmt_intl,
        "formato_nacional": fmt_nac, "codigo_pais": cod_pais, "pais_iso": pais_code,
        "pais_nombre": pais_nombre, "region": region, "operador": operador,
        "tipo_linea": tipo, "zonas_horarias": zonas, "valido": es_valido, "dorks": dorks,
    }
    guardar("osint_phone", resultado)
    return resultado

# ══════════════════════════════════════════
#   MÓDULO 4 — WHOIS DOMINIO
# ══════════════════════════════════════════
def osint_whois(dominio):
    separador("OSINT · WHOIS")
    resultado = {"dominio": dominio}

    if WHOIS_LIB_OK:
        try:
            w = whois_lib.whois(dominio)
            campos = ["registrar", "creation_date", "expiration_date",
                      "updated_date", "name_servers", "status", "emails", "org"]
            encontrado = False
            for k in campos:
                v = getattr(w, k, None)
                if v:
                    resultado[k] = str(v)
                    log("ok", f"{k:20}: {W}{str(v)[:60]}{RE}")
                    encontrado = True
            if encontrado:
                guardar("osint_whois", resultado)
                return resultado
        except Exception as e:
            log("warn", f"python-whois: {e}")

    ok, out = _correr_comando(["whois", dominio], timeout=15)
    if ok:
        for line in out.splitlines():
            ll = line.lower()
            if any(k in ll for k in ["registrar", "creation", "expir", "name server", "registrant", "email"]):
                log("ok", f"  {C}{line.strip()}{RE}")
        resultado["raw_disponible"] = True
    else:
        log("err", "whois no disponible (instala el cliente whois o python-whois)")
        if ES_WINDOWS:
            log("info", "En Windows: pip install python-whois")
        elif ES_TERMUX:
            log("info", "En Termux: pkg install whois")
        else:
            log("info", "En Linux: sudo apt install whois")
    guardar("osint_whois", resultado)
    return resultado

# ══════════════════════════════════════════
#   MÓDULO 5 — DNS LOOKUP
# ══════════════════════════════════════════
def dns_lookup(dominio):
    separador("DNS · LOOKUP")
    resultado = {"dominio": dominio, "registros": {}, "subdominios_comunes": []}
    tipos = {"A": "-type=A", "MX": "-type=MX", "NS": "-type=NS",
             "TXT": "-type=TXT", "CNAME": "-type=CNAME", "AAAA": "-type=AAAA"}

    for tipo, flag in tipos.items():
        ok, out = _correr_comando(["nslookup", flag, dominio])
        if ok:
            lineas = [l.strip() for l in out.splitlines()
                      if l.strip() and not l.startswith("Server") and not l.startswith("Address")]
            if lineas:
                resultado["registros"][tipo] = lineas
                log("ok", f"{tipo:6} → {lineas[0][:60]}")

    try:
        ips = socket.getaddrinfo(dominio, None)
        ipv4 = list({r[4][0] for r in ips if r[0].name == "AF_INET"})
        if ipv4:
            resultado["registros"]["A"] = ipv4
            log("ok", f"{'A':6} → {', '.join(ipv4)}")
    except Exception:
        pass

    log("info", "Para descubrimiento de subdominios usa herramientas pasivas")
    log("info", "como crt.sh o SecurityTrails, fuera del alcance de este módulo.")
    guardar("dns_lookup", resultado)
    return resultado

# ══════════════════════════════════════════
#   MÓDULO 6 — ANÁLISIS DE ARCHIVOS LOCALES
# ══════════════════════════════════════════
def archivo_hash(ruta):
    separador("ARCHIVO · HASHES")
    if not os.path.isfile(ruta):
        log("err", "Archivo no encontrado"); return None
    algos = {"MD5": hashlib.md5(), "SHA1": hashlib.sha1(), "SHA256": hashlib.sha256()}
    with open(ruta, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            for h in algos.values():
                h.update(chunk)
    resultado = {"archivo": ruta, "tamaño": os.path.getsize(ruta)}
    for nombre, h in algos.items():
        d = h.hexdigest()
        resultado[nombre] = d
        log("ok", f"{nombre:8}: {W}{d}{RE}")
    log("info", f"Tamaño   : {resultado['tamaño']} bytes")
    guardar("archivo_hash", resultado)
    return resultado

def archivo_info(ruta):
    separador("ARCHIVO · INFO")
    if not os.path.isfile(ruta):
        log("err", "Archivo no encontrado"); return None
    stat = os.stat(ruta)
    magics = {
        b'\x7fELF'     : "ELF (ejecutable Linux)",
        b'PK\x03\x04'  : "ZIP / APK / JAR",
        b'\xff\xd8\xff': "JPEG",
        b'\x89PNG'     : "PNG",
        b'%PDF'        : "PDF",
        b'MZ'          : "PE (ejecutable Windows)",
    }
    tipo = "Desconocido"
    with open(ruta, "rb") as f:
        header = f.read(8)
    for magic, nombre in magics.items():
        if header.startswith(magic):
            tipo = nombre; break
    info = {
        "nombre"    : os.path.basename(ruta),
        "ruta"      : os.path.abspath(ruta),
        "tipo"      : tipo,
        "tamaño"    : stat.st_size,
        "modificado": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "permisos"  : oct(stat.st_mode),
    }
    for k, v in info.items():
        log("ok", f"{k:12}: {W}{v}{RE}")
    guardar("archivo_info", info)
    return info

def archivo_strings(ruta, minlen=6):
    separador("ARCHIVO · STRINGS")
    if not os.path.isfile(ruta):
        log("err", "Archivo no encontrado"); return None
    strings = []
    current = ""
    with open(ruta, "rb") as f:
        data = f.read()
    for byte in data:
        c = chr(byte)
        if c.isprintable() and c != '\n':
            current += c
        else:
            if len(current) >= minlen:
                strings.append(current)
            current = ""
    if len(current) >= minlen:
        strings.append(current)
    log("ok", f"Strings encontrados: {len(strings)}")
    for s in strings[:40]:
        print(f"  {B}{s}{RE}")
    if len(strings) > 40:
        log("info", f"... y {len(strings)-40} más en el JSON exportado")
    datos_strings = {"archivo": ruta, "strings": strings}
    guardar("archivo_strings", datos_strings)
    return datos_strings

# ══════════════════════════════════════════
#   MÓDULO 7 — ANÁLISIS ESTÁTICO DE APK
# ══════════════════════════════════════════
def apk_info(ruta):
    separador("APK · ANÁLISIS ESTÁTICO")
    if not os.path.isfile(ruta):
        log("err", "Archivo no encontrado"); return None
    resultado = {
        "archivo": ruta, "tamaño": os.path.getsize(ruta),
        "hashes": {}, "permisos": [], "actividades": [], "strings_interes": []
    }
    for algo, h in [("md5", hashlib.md5()), ("sha256", hashlib.sha256())]:
        with open(ruta, "rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
        resultado["hashes"][algo] = h.hexdigest()
        log("ok", f"{algo.upper():8}: {W}{resultado['hashes'][algo]}{RE}")

    import zipfile
    try:
        with zipfile.ZipFile(ruta, "r") as z:
            nombres = z.namelist()
            log("ok", f"Archivos internos : {len(nombres)}")
            for n in nombres:
                if any(n.endswith(e) for e in [".dex", ".so", "MANIFEST.MF"]):
                    log("info", f"  → {n}")
            if "AndroidManifest.xml" in nombres:
                data = z.read("AndroidManifest.xml")
                strings = []
                cur = ""
                for byte in data:
                    c = chr(byte) if 32 <= byte < 127 else ""
                    if c:
                        cur += c
                    else:
                        if len(cur) >= 5:
                            strings.append(cur)
                        cur = ""
                if len(cur) >= 5:
                    strings.append(cur)
                perms   = sorted(set(s for s in strings if "permission" in s.lower()))[:20]
                activ   = sorted(set(s for s in strings if s.startswith("com.") or "activity" in s.lower()))[:20]
                interes = sorted(set(s for s in strings if any(
                    k in s.lower() for k in ["http", "password", "token", "key", "secret", "api", "login"]
                )))[:20]
                resultado["permisos"]        = perms
                resultado["actividades"]     = activ
                resultado["strings_interes"] = interes
                if perms:
                    log("warn", f"Permisos declarados ({len(perms)}):")
                    for p in perms[:8]:
                        print(f"  {Y}{p}{RE}")
                if interes:
                    log("warn", f"Strings de interés ({len(interes)}) — revisar manualmente:")
                    for s in interes[:8]:
                        print(f"  {R}{s}{RE}")
    except zipfile.BadZipFile:
        log("err", "APK corrupto o inválido")
    except Exception as e:
        log("err", str(e))

    ok, v = _correr_comando(["apktool", "--version"])
    if ok:
        log("ok", f"apktool {v.strip()} disponible → apktool d {ruta}")
    else:
        log("warn", "apktool no instalado (opcional, para descompilar)")
    guardar("apk_analisis", resultado)
    return resultado

# ══════════════════════════════════════════
#   MÓDULO 8 — INFO DE RED LOCAL 
# ══════════════════════════════════════════
def wifi_info():
    separador("RED · INFO LOCAL (este dispositivo)")
    resultado = {}
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip_local = s.getsockname()[0]
        s.close()
        resultado["ip_local"] = ip_local
        log("ok", f"IP local       : {W}{ip_local}{RE}")
    except Exception:
        log("err", "No se pudo obtener la IP local")

    try:
        resultado["hostname"] = socket.gethostname()
        log("ok", f"Hostname       : {W}{resultado['hostname']}{RE}")
    except Exception:
        pass

    if ES_WINDOWS:
        ok, out = _correr_comando(["ipconfig"])
        if ok:
            for line in out.splitlines():
                l = line.strip()
                if any(k in l for k in ["IPv4", "IPv6", "Adaptador", "Adapter"]):
                    print(f"  {B}{l}{RE}")
    else:
        for cmd in [["ip", "addr", "show"], ["ifconfig"]]:
            ok, out = _correr_comando(cmd)
            if ok:
                for line in out.splitlines():
                    l = line.strip()
                    if "inet " in l or "wlan" in l or "ether" in l:
                        print(f"  {B}{l}{RE}")
                break

    guardar("info_red_local", resultado)
    return resultado


# ══════════════════════════════════════════
#   MÓDULO 9 — BÚSQUEDA DE USERNAME
# ══════════════════════════════════════════
_PLATAFORMAS_USERNAME = {
    "GitHub"  : {"url": "https://github.com/{u}",          "metodo": "status", "no_existe_si": [404]},
    "GitLab"  : {"url": "https://gitlab.com/{u}",           "metodo": "status", "no_existe_si": [404]},
    "Reddit"  : {"url": "https://www.reddit.com/user/{u}/about.json", "metodo": "status", "no_existe_si": [404]},
    "Twitch"  : {"url": "https://www.twitch.tv/{u}",        "metodo": "texto", "patron_no_existe": ["sorry. unless you've got a time machine"]},
    "Steam"   : {"url": "https://steamcommunity.com/id/{u}", "metodo": "texto", "patron_no_existe": ["the specified profile could not be found"]},
    "Telegram": {"url": "https://t.me/{u}",                  "metodo": "texto", "patron_no_existe": ["if you have telegram, you can contact"], "invertir": True},
    "Keybase" : {"url": "https://keybase.io/{u}",            "metodo": "status", "no_existe_si": [404]},
    "Dev.to"  : {"url": "https://dev.to/{u}",                "metodo": "status", "no_existe_si": [404]},
}

def _check_username_en(nombre, plataforma, username, timeout):
    url = plataforma["url"].format(u=username)
    metodo = plataforma.get("metodo", "status")
    try:
        r = requests.get(
            url, timeout=timeout, allow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (compatible; AltKillZone-UsernameCheck/3.1)"}
        )
    except requests.exceptions.RequestException:
        return nombre, None, url, None

    if metodo == "status":
        existe = r.status_code not in plataforma["no_existe_si"]
        return nombre, existe, url, r.status_code

    if metodo == "texto":
        cuerpo = r.text.lower()
        encontrado_patron = any(p in cuerpo for p in plataforma["patron_no_existe"])
        if plataforma.get("invertir"):
            existe = encontrado_patron
        else:
            existe = not encontrado_patron
        return nombre, existe, url, r.status_code

    return nombre, None, url, r.status_code

def buscar_username(username, timeout=8, max_hilos=6):
    separador(f"BÚSQUEDA DE USERNAME → {username}")
    if not REQUESTS_OK:
        log("err", "Falta la librería requests → pip install requests"); return None

    log("info", f"Consultando {len(_PLATAFORMAS_USERNAME)} plataformas verificables...")
    print()
    resultado = {"username": username, "encontrados": [], "no_encontrados": [], "sin_verificar": []}

    with ThreadPoolExecutor(max_workers=max_hilos) as executor:
        futuros = [
            executor.submit(_check_username_en, nombre, datos, username, timeout)
            for nombre, datos in _PLATAFORMAS_USERNAME.items()
        ]
        filas = [fut.result() for fut in as_completed(futuros)]

    for nombre, existe, url, status in sorted(filas, key=lambda x: x[0]):
        if existe is True:
            log("ok", f"{nombre:<10} {G}✓ ENCONTRADO{RE}  {C}{url}{RE}")
            resultado["encontrados"].append({"plataforma": nombre, "url": url, "status": status})
        elif existe is False:
            log("info", f"{nombre:<10} {B}✗ no encontrado{RE}")
            resultado["no_encontrados"].append(nombre)
        else:
            log("warn", f"{nombre:<10} {Y}? no se pudo verificar (timeout/red){RE}")
            resultado["sin_verificar"].append(nombre)

    print()
    log("info", f"Encontrados: {len(resultado['encontrados'])} | "
                 f"No encontrados: {len(resultado['no_encontrados'])} | "
                 f"Sin verificar: {len(resultado['sin_verificar'])}")
    log("info", "Para más plataformas, verifica manualmente con los enlaces")
    log("info", "generados en el módulo de email (buscar_redes_sociales).")
    guardar("username_search", resultado)
    return resultado

# ══════════════════════════════════════════
#   MÓDULO 10 — VALIDADOR OFICIAL DE CURP 
# ══════════════════════════════════════════

def validar_curp_propio(curp):
    separador("CURP · VALIDACIÓN OFICIAL (propia)")
    curp = curp.strip().upper()
    patron = r'^[A-Z]{4}\d{6}[HM][A-Z]{5}[A-Z0-9]\d$'
    resultado = {"curp": curp, "formato_valido": bool(re.match(patron, curp))}

    if resultado["formato_valido"]:
        log("ok", "Formato de CURP válido (18 caracteres, estructura correcta).")
    else:
        log("warn", "El formato no coincide con la estructura estándar de CURP.")

    url_oficial = "https://www.gob.mx/curp/"
    resultado["url_validador_oficial"] = url_oficial
    log("info", f"Validador oficial (RENAPO): {C}{url_oficial}{RE}")
    log("warn", "Este módulo NO consulta ni expone datos de terceros.")
    log("warn", "Úsalo únicamente para verificar tu propio CURP en el sitio oficial.")

    try:
        abrir = input(f"  {Y}[?]{RE} ¿Abrir el validador oficial en el navegador? (s/n): ").strip().lower()
        if abrir == "s":
            webbrowser.open(url_oficial)
    except Exception:
        pass

    guardar("curp_validacion", resultado)
    return resultado

# ══════════════════════════════════════════
#   REPORTE CONSOLIDADO EN WORD
# ══════════════════════════════════════════
def _agregar_datos_recursivo(doc, datos):
    if isinstance(datos, dict):
        for clave, valor in datos.items():
            if isinstance(valor, (dict, list)) and valor:
                p = doc.add_paragraph()
                p.add_run(f"{clave}:").bold = True
                _agregar_datos_recursivo(doc, valor)
            else:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{clave}: ").bold = True
                p.add_run(str(valor))
    elif isinstance(datos, list):
        for item in datos:
            if isinstance(item, dict):
                texto = " | ".join(f"{k}: {v}" for k, v in item.items())
                doc.add_paragraph(texto, style="List Bullet")
            else:
                doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph(str(datos))


def generar_reporte_word(nombre_archivo=None):
    separador("REPORTE · GENERAR DOCUMENTO WORD")

    if not DOCX_OK:
        log("err", "Falta la librería python-docx → pip install python-docx")
        return None

    if not SESION_RESULTADOS:
        log("warn", "No hay investigaciones en esta sesión todavía.")
        return None

    doc = Document()

    doc.add_heading("ALT KILL ZONE — Reporte de Investigación OSINT", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    doc.add_paragraph("Autor de la herramienta: PARADOX MX")
    doc.add_paragraph(f"Total de investigaciones incluidas: {len(SESION_RESULTADOS)}")
    doc.add_paragraph(
        "Este reporte reúne información pública obtenida mediante técnicas de "
        "reconocimiento pasivo (OSINT). Su uso es responsabilidad exclusiva de "
        "quien lo generó."
    )
    doc.add_page_break()

    doc.add_heading("Contenido de la sesión", level=1)
    for i, r in enumerate(SESION_RESULTADOS, 1):
        doc.add_paragraph(f"{i}. [{r['tipo']}] {r['objetivo']}  —  {r['fecha']}", style="List Number")
    doc.add_page_break()

    for i, r in enumerate(SESION_RESULTADOS, 1):
        doc.add_heading(f"{i}. {r['tipo']} — {r['objetivo']}", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Fecha de investigación: {r['fecha']}").italic = True
        doc.add_paragraph()
        _agregar_datos_recursivo(doc, r["datos"])
        doc.add_page_break()

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre = nombre_archivo or f"akz_reporte_{ts}.docx"
    directorio = _directorio_salida()
    ruta = os.path.join(directorio, nombre)
    try:
        doc.save(ruta)
        log("ok", f"Reporte Word generado → {W}{ruta}{RE}")
        return ruta
    except Exception:
        try:
            ruta = nombre
            doc.save(ruta)
            log("ok", f"Reporte Word generado → {W}{ruta}{RE}")
            return ruta
        except Exception as e2:
            log("err", f"No se pudo guardar el reporte Word: {e2}")
            return None

# ══════════════════════════════════════════
#   MENÚ INTERACTIVO PRINCIPAL
# ══════════════════════════════════════════
def _pedir(texto):
    while True:
        valor = input(f"  {Y}[?]{RE} {texto}: ").strip()
        if valor:
            return valor
        log("err", "No puedes dejarlo vacío.")

def menu_principal():
    while True:
        show_banner()
        n_sesion = len(SESION_RESULTADOS)
        print(f"""
{R}╔══════════════════════════════════════════════════════╗
║              🧭  MENÚ PRINCIPAL — ALT KILL ZONE        ║
╠══════════════════════════════════════════════════════╣
║ {G}[1]{W}  🌐 OSINT · IP                                    {R}║
║ {G}[2]{W}  📧 OSINT · Email                                 {R}║
║ {G}[3]{W}  📱 OSINT · Teléfono                               {R}║
║ {G}[4]{W}  🗂  WHOIS · Dominio                                {R}║
║ {G}[5]{W}  🧬 DNS · Lookup                                   {R}║
║ {G}[6]{W}  🗄  Archivo · Hashes                               {R}║
║ {G}[7]{W}  📄 Archivo · Info                                  {R}║
║ {G}[8]{W}  🔡 Archivo · Strings                               {R}║
║ {G}[9]{W}  📦 APK · Análisis estático                         {R}║
║ {G}[10]{W} 📶 Red local · Info                                {R}║
║ {G}[11]{W} 🕵  Username · Búsqueda multiplataforma            {R}║
║ {G}[12]{W} 🪪 CURP · Validador oficial (propio, gob.mx)        {R}║
║ {G}[13]{W} 📝 Generar reporte Word ({n_sesion} en sesión)                 {R}║
║ {G}[0]{W}  ❌ Salir                                           {R}║
╚══════════════════════════════════════════════════════╝{RE}
""")
        op = input(f"  {Y}[?]{RE} Selecciona una opción (0-13): ").strip()

        if op == "1":
            ip = _pedir("IP a investigar (ej: 8.8.8.8)")
            datos = osint_ip(ip)
            _agregar_sesion("OSINT IP", ip, datos)

        elif op == "2":
            email = _pedir("Email a investigar")
            obj = OsintEmail(email)
            obj.menu()
            _agregar_sesion("OSINT Email", email, obj.investigacion)

        elif op == "3":
            numero = _pedir("Número telefónico (ej: +5212345678901)")
            datos = osint_phone(numero)
            _agregar_sesion("OSINT Teléfono", numero, datos)

        elif op == "4":
            dominio = _pedir("Dominio a consultar (ej: ejemplo.com)")
            datos = osint_whois(dominio)
            _agregar_sesion("WHOIS", dominio, datos)

        elif op == "5":
            dominio = _pedir("Dominio para DNS lookup")
            datos = dns_lookup(dominio)
            _agregar_sesion("DNS Lookup", dominio, datos)

        elif op == "6":
            ruta = _pedir("Ruta del archivo")
            datos = archivo_hash(ruta)
            _agregar_sesion("Archivo · Hashes", ruta, datos)

        elif op == "7":
            ruta = _pedir("Ruta del archivo")
            datos = archivo_info(ruta)
            _agregar_sesion("Archivo · Info", ruta, datos)

        elif op == "8":
            ruta = _pedir("Ruta del archivo")
            datos = archivo_strings(ruta)
            _agregar_sesion("Archivo · Strings", ruta, datos)

        elif op == "9":
            ruta = _pedir("Ruta del archivo APK")
            datos = apk_info(ruta)
            _agregar_sesion("APK", ruta, datos)

        elif op == "10":
            datos = wifi_info()
            _agregar_sesion("Red local", "este dispositivo", datos)

        elif op == "11":
            username = _pedir("Username a buscar")
            try:
                timeout_str = input(f"  {Y}[?]{RE} Timeout en segundos (Enter = 8): ").strip()
                timeout = int(timeout_str) if timeout_str else 8
            except ValueError:
                timeout = 8
            datos = buscar_username(username, timeout=timeout)
            _agregar_sesion("Username", username, datos)

        elif op == "12":
            curp = _pedir("Tu CURP (18 caracteres)")
            datos = validar_curp_propio(curp)
            _agregar_sesion("CURP (validación propia)", curp, datos)

        elif op == "13":
            generar_reporte_word()

        elif op == "0":
            if SESION_RESULTADOS:
                resp = input(
                    f"  {Y}[?]{RE} Tienes {len(SESION_RESULTADOS)} investigación(es) sin exportar. "
                    f"¿Generar reporte Word antes de salir? (s/n): "
                ).strip().lower()
                if resp == "s":
                    generar_reporte_word()
            log("info", "Saliendo... ¡Hasta la próxima!")
            break

        else:
            log("err", "Opción inválida")

        if op != "0":
            input(f"\n  {C}[↵] Presiona Enter para volver al menú principal...{RE}")
            limpiar_pantalla()

# ══════════════════════════════════════════
#   MAIN / CLI
# ══════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(
        description="ALT KILL ZONE v3.1 | PARADOX MX — OSINT y reconocimiento pasivo",
        formatter_class=argparse.RawTextHelpFormatter,
        epilog="""
Ejemplos:
  %(prog)s                          (sin argumentos = abre el menú interactivo)
  %(prog)s osint --ip 8.8.8.8
  %(prog)s osint --email usuario@dominio.com
  %(prog)s osint --phone +5212345678901
  %(prog)s osint --whois ejemplo.com
  %(prog)s dns --lookup ejemplo.com
  %(prog)s archivo --hash ruta/al/archivo
  %(prog)s apk --info app.apk
  %(prog)s red --info
  %(prog)s username --buscar mi_usuario
  %(prog)s curp --validar TUCURP180101HDFXXX01
        """
    )
    sub = parser.add_subparsers(dest="modulo")

    p = sub.add_parser("osint", help="IP / email / teléfono / WHOIS")
    p.add_argument("--ip",    metavar="IP",     help="Geolocalización de IP")
    p.add_argument("--email", metavar="EMAIL",  help="Análisis de email")
    p.add_argument("--phone", metavar="NÚMERO", help="Info de número (ej: +521234567890)")
    p.add_argument("--whois", metavar="DOMINIO", help="WHOIS de dominio")

    p = sub.add_parser("dns", help="DNS Lookup de un dominio")
    p.add_argument("--lookup", metavar="DOMINIO")

    p = sub.add_parser("archivo", help="Análisis de archivos locales")
    p.add_argument("--hash",    metavar="RUTA")
    p.add_argument("--info",    metavar="RUTA")
    p.add_argument("--strings", metavar="RUTA")

    p = sub.add_parser("apk", help="Análisis estático de APK")
    p.add_argument("--info", metavar="RUTA")

    p = sub.add_parser("red", help="Info de red local (este dispositivo)")
    p.add_argument("--info", action="store_true", help="Mostrar IP/hostname/interfaces locales")

    p = sub.add_parser("username", help="Buscar username en plataformas públicas")
    p.add_argument("--buscar",  metavar="USERNAME", help="Username a buscar")
    p.add_argument("--timeout", metavar="SEG", type=int, default=6)

    p = sub.add_parser("curp", help="Validador oficial de CURP (propio)")
    p.add_argument("--validar", metavar="CURP", help="CURP propio a validar")

    p = sub.add_parser("reporte", help="Generar reporte Word de la sesión actual")

    p = sub.add_parser("menu", help="Abrir el menú interactivo")

    args = parser.parse_args()

    if not args.modulo or args.modulo == "menu":
        menu_principal()
        return

    show_banner()

    if args.modulo == "osint":
        if args.ip:    osint_ip(args.ip)
        if args.email: osint_email(args.email)
        if args.phone: osint_phone(args.phone)
        if args.whois: osint_whois(args.whois)

    elif args.modulo == "dns":
        if args.lookup: dns_lookup(args.lookup)

    elif args.modulo == "archivo":
        if args.hash:    archivo_hash(args.hash)
        if args.info:    archivo_info(args.info)
        if args.strings: archivo_strings(args.strings)

    elif args.modulo == "apk":
        if args.info: apk_info(args.info)

    elif args.modulo == "red":
        if args.info: wifi_info()

    elif args.modulo == "username":
        if args.buscar: buscar_username(args.buscar, timeout=args.timeout)

    elif args.modulo == "curp":
        if args.validar: validar_curp_propio(args.validar)

    elif args.modulo == "reporte":
        generar_reporte_word()


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print(f"\n\n[!] Interrumpido por el usuario. Saliendo...")
        sys.exit(0)
    except Exception as e:
        print(f"\n[x] Error crítico: {e}")
        sys.exit(1)